from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


def add_text(document, text, bold=False, size=14, italic=False):
    p = document.add_paragraph()
    run = p.add_run(text)
    run.italic = italic
    run.bold = bold
    run.font.size = Pt(size)

def add_title(document, text, bold=True, size=16, italic=False):
    p = document.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)

def make_CV(personal_data, purpose, exp, education, personal_qualities, user_id: int):
    doc = Document()

    title = doc.add_heading('Резюме', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_title(doc, 'Личная информация')
    add_text(doc, personal_data)

    doc.add_paragraph()
    add_title(doc, 'Цель резюме')
    add_text(doc, purpose)

    doc.add_paragraph()
    add_title(doc, 'Профессиональный опыт')
    add_text(doc, exp)

    doc.add_paragraph()
    add_title(doc, 'Образование')
    add_text(doc, education)

    doc.add_paragraph()
    add_title(doc, 'Личные качества')
    add_text(doc, personal_qualities)

    file_path = f'data/{user_id}.docx'

    doc.save(file_path)
    return file_path
